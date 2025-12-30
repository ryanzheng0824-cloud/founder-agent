import streamlit as st
from tavily import TavilyClient
from openai import OpenAI  # <--- æ ¸å¿ƒå˜åŒ–ï¼šç”¨ OpenAI åº“è¿æ¥ DeepSeek
import os
from dotenv import load_dotenv

# --- 1. å…¨å±€é…ç½® ---
st.set_page_config(page_title="é›„å¿ƒèŸÂ·åˆ›ä¸šå‚è°‹", page_icon="ğŸ¦", layout="wide")
# ... (st.set_page_config ä»£ç åœ¨ä¸Šé¢) ...

# === ğŸ¦ éšè—ç•Œé¢å…ƒç´ çš„ CSS é»‘ç§‘æŠ€ ===
hide_streamlit_style = """
<style>
/* éšè—å³ä¸Šè§’çš„ä¸‰é“æ èœå• */
#MainMenu {visibility: hidden;}
/* éšè—åº•éƒ¨çš„ "Made with Streamlit" */
footer {visibility: hidden;}
/* éšè—é¡¶éƒ¨çš„å¯¼èˆªæ çº¿ */
header {visibility: hidden;}
</style>
"""
st.markdown(hide_streamlit_style, unsafe_allow_html=True)

# ... (å‰©ä¸‹çš„ä»£ç åœ¨ä¸‹é¢) ...
load_dotenv()

# è·å– Key
deepseek_key = os.environ.get("DEEPSEEK_API_KEY")
tavily_key = os.environ.get("TAVILY_API_KEY")

# æ£€æŸ¥ Key æ˜¯å¦å­˜åœ¨
if not deepseek_key:
    st.error("âŒ æœªæ‰¾åˆ° DEEPSEEK_API_KEYï¼Œè¯·æ£€æŸ¥ .env æ–‡ä»¶")
    st.stop()

# åˆå§‹åŒ– DeepSeek å®¢æˆ·ç«¯
client = OpenAI(
    api_key=deepseek_key, 
    base_url="https://api.deepseek.com"  # <--- æŒ‡å‘ DeepSeek æœåŠ¡å™¨
)

# --- 2. ç•Œé¢è®¾è®¡ ---
st.title("ğŸ¦ é›„å¿ƒèŸÂ·è¶…çº§è½»åˆ›ä¸ªä½“å‚è°‹")
st.markdown("### ğŸš€ ä¸“ä¸ºé›„å¿ƒèŸä¼šå‘˜æ‰“é€  | æ·±åº¦å•†ä¸šè¯„ä¼°ç³»ç»Ÿ")

# ä¾§è¾¹æ è¯´æ˜
with st.sidebar:
    st.info("ğŸ’¡ **ä½¿ç”¨æŒ‡å—**")
    st.markdown("""
    è¾“å…¥ä½ æƒ³åšçš„é¡¹ç›®ï¼ˆå¦‚ï¼š*é²œèŠ±åº—ã€æ”¶çº³å¸ˆã€è½¦ä½æŠ•èµ„*ï¼‰ï¼ŒAI å°†ä¸ºä½ ï¼š
    1. ğŸ•µï¸â€â™‚ï¸ **å…¨ç½‘è°ƒç ”**ï¼šæœå¯»æœ€æ–°è¡Œæƒ…
    2. ğŸ§  **æ·±åº¦è¯„ä¼°**ï¼šDeepSeek æ‹†è§£æ¨¡å¼
    3. ğŸ’° **ç®—è´¦é¿å‘**ï¼šè®¡ç®—å›æœ¬å‘¨æœŸ
    """)
    st.divider()
    st.caption("Powered by DeepSeek V3 & Tavily")

# --- 3. æ ¸å¿ƒé€»è¾‘ ---
topic = st.text_input("ğŸ‘‡ è¯·è¾“å…¥ä½ æƒ³è¯„ä¼°çš„åˆ›ä¸šé¡¹ç›®:", placeholder="ä¾‹å¦‚ï¼šåœ¨ä¸‰çº¿åŸå¸‚å¼€ä¸€å®¶è‡ªåŠ©æ´—è½¦åº—")
start_btn = st.button("ğŸš€ å¼€å§‹æ·±åº¦è¯„ä¼°", type="primary")

if start_btn and topic:
    # çŠ¶æ€å®¹å™¨
    with st.status("âš™ï¸ å‚è°‹æ­£åœ¨å·¥ä½œä¸­...", expanded=True) as s:
        
        # æ­¥éª¤ A: è”ç½‘æœç´¢
        s.write("ğŸ•µï¸â€â™‚ï¸ æ­£åœ¨å…¨ç½‘æœé›†æƒ…æŠ¥ (Tavily)...")
        tavily = TavilyClient(api_key=tavily_key)
        
        # è‡ªåŠ¨æŠŠç”¨æˆ·çš„è¯æ‰©å±•ï¼Œæœå¾—æ›´ç»†
        # åŠ ä¸Šå¹´ä»½å’Œâ€œæœ€æ–°â€å…³é”®è¯ï¼Œå¼ºåˆ¶æœç´¢å¼•æ“æ‰¾è¿‘æœŸçš„
        import datetime
        current_year = datetime.datetime.now().year
        # è·å–ä»Šå¤©çš„å…·ä½“æ—¥æœŸï¼Œä¾‹å¦‚ "2025å¹´12æœˆ30æ—¥"
        today_str = datetime.datetime.now().strftime("%Yå¹´%mæœˆ%dæ—¥")
        
        # é’ˆå¯¹é‡‘è/æ—¶æ•ˆæ€§å¼ºçš„å…³é”®è¯ï¼Œå¼ºåˆ¶åŠ ä¸Šæ—¥æœŸå’Œâ€œæœ€æ–°â€
        # æœç´¢è¯å˜æˆï¼š "ç™½é“¶ ä»·æ ¼èµ°åŠ¿ 2025å¹´12æœˆ30æ—¥ æœ€æ–°è¡Œæƒ… è¿‘ä¸€å‘¨æ¶¨è·ŒåŸå› "
        search_query = f"{topic} ä»·æ ¼èµ°åŠ¿ {today_str} æœ€æ–°è¡Œæƒ… è¿‘ä¸€å‘¨æ¶¨è·ŒåŸå› "
        
        s.write(f"ğŸ•µï¸â€â™‚ï¸ æ­£åœ¨æœé›† {today_str} çš„æœ€æ–°æƒ…æŠ¥...") # æç¤ºè¯­ä¹Ÿæ”¹ä¸€ä¸‹ï¼Œçœ‹ç€æ›´çˆ½
        
        tavily = TavilyClient(api_key=tavily_key)
        search_res = tavily.search(query=search_query, search_depth="advanced", max_results=5)        
        # æ•´ç†æœç´¢ç»“æœ
        context = "\n".join([f"ã€æ¥æºï¼š{r['title']}ã€‘{r['content']}" for r in search_res['results']])
        s.write("âœ… æƒ…æŠ¥æœé›†å®Œæ¯•ï¼")
        
        # æ­¥éª¤ B: DeepSeek æ€è€ƒ
        s.write("ğŸ§  DeepSeek V3 æ­£åœ¨æ·±åº¦åˆ†æå•†ä¸šæ¨¡å¼...")
        
        # ä½ çš„â€œæ¯’èˆŒå¯¼å¸ˆâ€ Prompt
        prompt = f"""
        ã€èº«ä»½ã€‘ä½ æ˜¯ä¸€ä½æ‹¥æœ‰ 20 å¹´å®æˆ˜ç»éªŒçš„èµ„æ·±åˆ›ä¸šå¯¼å¸ˆï¼Œä¸“ä¸ºâ€œé›„å¿ƒèŸâ€ä¸ªä½“åˆ›ä¸šè€…æœåŠ¡ã€‚ä½ ç†Ÿæ‚‰ä¸­å›½ä¸‹æ²‰å¸‚åœºã€å®ä½“åº—é€»è¾‘å’Œç”µå•†ç©æ³•ã€‚
        
        ã€ç”¨æˆ·æƒ³åšã€‘"{topic}"
        
        ã€å…¨ç½‘æƒ…æŠ¥ã€‘
        {context}
        
        ã€ä»»åŠ¡ã€‘è¯·åŸºäºæƒ…æŠ¥ï¼Œæ’°å†™ã€Šé¡¹ç›®å¯è¡Œæ€§æ·±åº¦è¯„æµ‹ã€‹ã€‚
        
        ã€è¦æ±‚ã€‘
        1. **æ‹’ç»åºŸè¯**ï¼šç”¨æ•°æ®è¯´è¯ï¼ŒçŠ€åˆ©ç‚¹è¯„ï¼Œä¸è®²æ­£ç¡®çš„åºŸè¯ã€‚
        2. **å¿…é¡»åŒ…å«ä»¥ä¸‹æ¨¡å—**ï¼š
           - ğŸ“Š **å¸‚åœºçº¢è“æµ·**ï¼šç”¨æ•°æ®åˆ¤æ–­é¥±å’Œåº¦ã€‚
           - ğŸ’° **ç®—ç¬”è´¦**ï¼šé¢„ä¼°å®¢å•ä»·ã€æ¯›åˆ©ã€ç›ˆäºå¹³è¡¡ç‚¹ã€å›æœ¬å‘¨æœŸï¼ˆå¿…é¡»ç»™å‡ºä¼°ç®—æ•°å­—ï¼‰ã€‚
           - ğŸšš **è¿›è´§å®æ“**ï¼šç»™å‡ºå…·ä½“çš„å¹³å°åç§°ï¼ˆå¦‚1688å…³é”®è¯ï¼‰ã€æ‰¹å‘å¸‚åœºåå­—æˆ–APPã€‚
           - âš ï¸ **åŠé€€æŒ‡å—**ï¼šç›´å‡»ç—›ç‚¹ï¼Œä»€ä¹ˆæ ·çš„äººåƒä¸‡åˆ«å¹²è¿™ä¸ªã€‚
        3. **ç»“å°¾æ¨è**ï¼šç»™å‡º 0-10 åˆ†çš„æ¨èæŒ‡æ•°ï¼Œå¹¶ä¸€å¥è¯æ€»ç»“ã€‚
        4. **æ ¼å¼**ï¼šä½¿ç”¨ Markdownï¼Œæ’ç‰ˆæ¸…æ™°ï¼Œå¤šç”¨ Emojiã€‚
        """
        
        # å‘é€ç»™ DeepSeek
        response = client.chat.completions.create(
            model="deepseek-chat",  # æŒ‡å®šæ¨¡å‹
            messages=[
                {"role": "system", "content": "ä½ æ˜¯ä¸€ä¸ªä¸“ä¸šã€çŠ€åˆ©ã€æ•°æ®é©±åŠ¨çš„å•†ä¸šåˆ†æå¸ˆã€‚"},
                {"role": "user", "content": prompt},
            ],
            stream=False
        )
        
        # è·å–ç»“æœ
        article = response.choices[0].message.content
        s.update(label="âœ… è¯„ä¼°æŠ¥å‘Šå·²ç”Ÿæˆ", state="complete", expanded=False)

# ... (å‰é¢çš„ä»£ç ä¸å˜) ...
    
    # ... (å‰é¢çš„ä»£ç ä¸å˜) ...
    
    # --- 4. ç»“æœå±•ç¤º ---
    st.divider()
    st.markdown(article)

    # === ğŸŒŸ å‡çº§ç‰ˆï¼šç”Ÿæˆ Word æ–‡æ¡£ ===
    from docx import Document
    from io import BytesIO

    # åˆ›å»ºä¸€ä¸ªå†…å­˜é‡Œçš„ Word æ–‡æ¡£
    doc = Document()
    doc.add_heading(f'ğŸ¦ é›„å¿ƒèŸÂ·åˆ›ä¸šè¯„æµ‹ï¼š{topic}', 0)
    
    # æŠŠç”Ÿæˆçš„æŠ¥å‘Šå†™å…¥ Word (æ³¨æ„ï¼šWord ä¸ä¼šè‡ªåŠ¨æ¸²æŸ“ Markdown çš„åŠ ç²—æ ¼å¼ï¼Œä½†å†…å®¹éƒ½åœ¨)
    doc.add_paragraph(article)
    doc.add_paragraph('\n\n(ç”± DeepSeek & é›„å¿ƒèŸ AI å‚è°‹ç”Ÿæˆ)')

    # ä¿å­˜åˆ°å†…å­˜
    binary_output = BytesIO()
    doc.save(binary_output)
    binary_output.seek(0)
    
    # æ–‡ä»¶å
    file_name = f"åˆ›ä¸šè¯„æµ‹_{topic}_{datetime.datetime.now().strftime('%Y%m%d')}.docx"
    
    st.download_button(
        label="ğŸ“¥ ä¸‹è½½ Word æŠ¥å‘Š (æ‰‹æœºå‹å¥½ç‰ˆ)",
        data=binary_output,
        file_name=file_name,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    # ==============================
    
    # ... (åº•éƒ¨ç‰ˆæƒä¸å˜) ...
    # ==============================

    # åº•éƒ¨ç‰ˆæƒ
    st.divider()
    st.caption("ğŸ¦ é›„å¿ƒèŸå†…éƒ¨å·¥å…· | æ•°æ®ä»…ä¾›å‚è€ƒï¼ŒæŠ•èµ„éœ€è°¨æ…")
